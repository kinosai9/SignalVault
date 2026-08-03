"""M3-C-2a: DOCX 文件支持测试。

覆盖：
- file_content_extractor._extract_docx 段落+表格提取与质量评估
- file_profile.profile_uploaded_file 对 .docx 的 supported/encoding 判定
- is_docx_file / ALL_ACCEPTED_EXTENSIONS 含 .docx
"""

from __future__ import annotations

from pathlib import Path


def _make_docx(
    path: Path,
    paragraphs: list[str] | None = None,
    table: list[list[str]] | None = None,
) -> Path:
    """用 python-docx 创建一个测试用 .docx。

    table: 单个表格，list of rows，每行是 list of cell 字符串。
    """
    from docx import Document

    doc = Document()
    for p in paragraphs or []:
        doc.add_paragraph(p)
    if table:
        t = doc.add_table(rows=len(table), cols=len(table[0]))
        for i, row in enumerate(table):
            for j, cell in enumerate(row):
                t.cell(i, j).text = cell
    doc.save(str(path))
    return path


# ── file_content_extractor._extract_docx ──────────────────────────────────────


def test_extract_docx_paragraphs_and_tables(tmp_path: Path) -> None:
    """DOCX 段落与表格文本都被提取。"""
    from signalvault.sources.file_content_extractor import _extract_docx

    docx_path = _make_docx(
        tmp_path / "report.docx",
        paragraphs=[
            "宁德时代投资观点",
            "储能需求持续增长，看好长期逻辑。" * 12,
        ],
        table=[["指标", "数值"], ["营收增速", "35%"]],
    )
    result = _extract_docx(docx_path, "hash123", ".docx", "report.docx")
    assert result.extension == ".docx"
    assert "宁德时代投资观点" in result.text
    assert "储能需求持续增长" in result.text
    assert "营收增速" in result.text  # 表格内容
    assert result.parse_quality == "good"
    assert result.encoding == "binary"
    assert result.blocks_count >= 3


def test_extract_docx_short_content_minimal(tmp_path: Path) -> None:
    """内容过短时 parse_quality 为 minimal。"""
    from signalvault.sources.file_content_extractor import _extract_docx

    docx_path = _make_docx(tmp_path / "short.docx", paragraphs=["短"])
    result = _extract_docx(docx_path, "h", ".docx", "short.docx")
    assert result.parse_quality == "minimal"
    assert result.quality_warnings


def test_extract_docx_title_from_first_paragraph(tmp_path: Path) -> None:
    """无显式标题时，取首个非空段落作为 title。"""
    from signalvault.sources.file_content_extractor import _extract_docx

    docx_path = _make_docx(
        tmp_path / "t.docx",
        paragraphs=["这是标题", "正文内容" * 50],
    )
    result = _extract_docx(docx_path, "h", ".docx", "t.docx")
    assert result.title == "这是标题"


# ── file_profile.profile_uploaded_file ────────────────────────────────────────


def test_profile_docx_supported(tmp_path: Path) -> None:
    """profile_uploaded_file 对 .docx 返回 supported=True，encoding=binary。"""
    from signalvault.sources.file_profile import profile_uploaded_file

    docx_path = _make_docx(
        tmp_path / "ok.docx",
        paragraphs=["投资研究文档", "详细分析内容。" * 30],
    )
    profile = profile_uploaded_file(docx_path, "ok.docx")
    assert profile.supported is True
    assert profile.extension == ".docx"
    assert profile.detected_encoding == "binary"
    assert profile.content_hash  # 已计算
    assert profile.parse_quality in ("good", "degraded")


def test_profile_docx_corrupt_unsupported(tmp_path: Path) -> None:
    """损坏的 .docx（非 zip）返回 supported=False。"""
    from signalvault.sources.file_profile import profile_uploaded_file

    bad = tmp_path / "broken.docx"
    bad.write_bytes(b"not a real docx file content")
    profile = profile_uploaded_file(bad, "broken.docx")
    assert profile.supported is False
    assert "DOCX" in (profile.unsupported_reason or "")


def test_profile_docx_text_extracted_length(tmp_path: Path) -> None:
    """profile 的 extracted_text_length 反映提取的文本长度。"""
    from signalvault.sources.file_profile import profile_uploaded_file

    docx_path = _make_docx(
        tmp_path / "len.docx",
        paragraphs=["段落一", "段落二内容较长。" * 20],
    )
    profile = profile_uploaded_file(docx_path, "len.docx")
    assert profile.supported is True
    assert profile.extracted_text_length > 0
    assert profile.extracted_blocks_count >= 2


# ── 扩展名白名单 ──────────────────────────────────────────────────────────────


def test_docx_in_accepted_extensions() -> None:
    """M3-C-2a: .docx 在 ALL_ACCEPTED_EXTENSIONS 中。"""
    from signalvault.sources.file_profile import (
        ALL_ACCEPTED_EXTENSIONS,
        ALLOWED_DOC_EXTENSIONS,
        is_docx_file,
    )

    assert ".docx" in ALL_ACCEPTED_EXTENSIONS
    assert ".docx" in ALLOWED_DOC_EXTENSIONS
    assert is_docx_file("report.docx") is True
    assert is_docx_file("report.pdf") is False
    assert is_docx_file("report.txt") is False
